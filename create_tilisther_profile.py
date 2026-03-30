#!/usr/bin/env python3
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

# Create document
doc = Document()

# Set up styles
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# Title
title = doc.add_heading('TILISTHER CONSTRUCTION COMPANY', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in title.runs:
    run.font.color.rgb = RGBColor(204, 85, 0)  # Orange brand color

# Subtitle
subtitle = doc.add_paragraph('Building Excellence, Delivering Quality')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in subtitle.runs:
    run.font.size = Pt(14)
    run.font.italic = True
    run.font.color.rgb = RGBColor(100, 100, 100)

doc.add_paragraph()

# Company Overview
doc.add_heading('COMPANY OVERVIEW', level=1)
doc.add_paragraph(
    'Tilisther Construction Company is a leading construction firm based in Kenya, '
    'specializing in residential, commercial, and infrastructure development projects. '
    'With a commitment to quality craftsmanship, innovative solutions, and timely delivery, '
    'we have established ourselves as a trusted partner in the construction industry.'
)

doc.add_paragraph(
    'Founded on the principles of integrity, professionalism, and excellence, '
    'Tilisther Construction has successfully completed numerous projects across Kenya, '
    'ranging from private residences to large-scale commercial developments.'
)

# Vision & Mission
doc.add_heading('OUR VISION', level=1)
doc.add_paragraph(
    'To be the most trusted and innovative construction company in East Africa, '
    'setting the standard for quality, sustainability, and client satisfaction.'
)

doc.add_heading('OUR MISSION', level=1)
doc.add_paragraph(
    'To deliver exceptional construction services through skilled craftsmanship, '
    'modern technology, and sustainable practices while building lasting relationships '
    'with our clients, partners, and communities.'
)

# Core Values
doc.add_heading('CORE VALUES', level=1)
values = [
    ('Quality Excellence', 'We never compromise on the quality of materials and workmanship.'),
    ('Integrity', 'We conduct our business with honesty, transparency, and ethical practices.'),
    ('Innovation', 'We embrace modern construction techniques and technologies.'),
    ('Client Focus', 'We prioritize our clients\' needs and exceed their expectations.'),
    ('Safety', 'We maintain the highest safety standards on all our project sites.'),
    ('Sustainability', 'We are committed to environmentally responsible construction practices.')
]

for value, description in values:
    p = doc.add_paragraph()
    p.add_run(f'• {value}: ').bold = True
    p.add_run(description)

# Services
doc.add_heading('OUR SERVICES', level=1)

services = [
    ('Residential Construction', 'Custom homes, apartment complexes, and housing developments built to the highest standards.'),
    ('Commercial Buildings', 'Office complexes, shopping malls, hotels, and retail spaces designed for functionality and aesthetics.'),
    ('Industrial Construction', 'Factories, warehouses, and manufacturing facilities with specialized requirements.'),
    ('Infrastructure Development', 'Roads, bridges, drainage systems, and public works projects.'),
    ('Renovation & Remodeling', 'Building upgrades, expansions, and modernization projects.'),
    ('Project Management', 'End-to-end project planning, coordination, and supervision.')
]

for service, desc in services:
    p = doc.add_paragraph()
    p.add_run(f'• {service}\n').bold = True
    p.add_run(desc)
    doc.add_paragraph()

# Why Choose Us
doc.add_heading('WHY CHOOSE TILISTHER CONSTRUCTION?', level=1)

reasons = [
    'Experienced team of engineers, architects, and construction professionals',
    'Proven track record of successful project completion',
    'Competitive pricing without compromising quality',
    'Timely project delivery',
    'Comprehensive warranty and after-sales support',
    'Use of high-quality, locally sourced materials',
    'Full compliance with Kenyan building codes and regulations',
    'Strong relationships with trusted suppliers and subcontractors'
]

for reason in reasons:
    doc.add_paragraph(f'✓ {reason}')

# Projects
doc.add_heading('NOTABLE PROJECTS', level=1)
doc.add_paragraph(
    'Tilisther Construction has successfully delivered projects across various sectors. '
    'Our portfolio includes:'
)

projects = [
    'Residential complexes in Nairobi and surrounding areas',
    'Commercial office buildings in major business districts',
    'Educational institutions and schools',
    'Healthcare facilities and clinics',
    'Retail spaces and shopping centers',
    'Infrastructure projects for local governments'
]

for project in projects:
    doc.add_paragraph(f'• {project}')

# Team
doc.add_heading('OUR TEAM', level=1)
doc.add_paragraph(
    'Our team comprises highly skilled professionals including:'
)

team = [
    'Project Managers',
    'Civil Engineers',
    'Architects',
    'Quantity Surveyors',
    'Site Supervisors',
    'Skilled Tradespeople (Masons, Carpenters, Electricians, Plumbers)',
    'Health and Safety Officers'
]

for member in team:
    doc.add_paragraph(f'• {member}')

# Quality Assurance
doc.add_heading('QUALITY ASSURANCE', level=1)
doc.add_paragraph(
    'At Tilisther Construction, quality is not just a goal—it\'s our standard. '
    'We implement rigorous quality control measures at every stage of construction, from '
    'material selection to final inspection. Our quality assurance process includes:'
)

qa_items = [
    'Regular site inspections and progress monitoring',
    'Material testing and certification',
    'Compliance with international building standards',
    'Third-party quality audits',
    'Client walkthroughs and approval milestones'
]

for item in qa_items:
    doc.add_paragraph(f'• {item}')

# Safety
doc.add_heading('HEALTH & SAFETY', level=1)
doc.add_paragraph(
    'We prioritize the safety of our workers, clients, and the public. '
    'Our comprehensive Health and Safety Policy ensures:'
)

safety_items = [
    'Mandatory safety training for all personnel',
    'Regular safety audits and risk assessments',
    'Proper personal protective equipment (PPE) provision',
    'Emergency response protocols',
    'Compliance with Occupational Safety and Health Administration (OSHA) standards'
]

for item in safety_items:
    doc.add_paragraph(f'• {item}')

# Sustainability
doc.add_heading('SUSTAINABILITY COMMITMENT', level=1)
doc.add_paragraph(
    'Tilisther Construction is committed to sustainable building practices. '
    'We incorporate environmentally friendly methods and materials wherever possible, including:'
)

sustainability_items = [
    'Energy-efficient building designs',
    'Water conservation systems',
    'Waste reduction and recycling programs',
    'Use of sustainable and locally sourced materials',
    'Green building certifications where applicable'
]

for item in sustainability_items:
    doc.add_paragraph(f'• {item}')

# Contact Information
doc.add_heading('CONTACT US', level=1)

contact_info = doc.add_paragraph()
contact_info.add_run('TILISTHER CONSTRUCTION COMPANY\n').bold = True
contact_info.add_run('Nairobi, Kenya\n\n')
contact_info.add_run('Phone: ').bold = True
contact_info.add_run('[Phone Number]\n')
contact_info.add_run('Email: ').bold = True
contact_info.add_run('[Email Address]\n')
contact_info.add_run('Website: ').bold = True
contact_info.add_run('[Website URL]\n')

# Footer
doc.add_paragraph()
footer = doc.add_paragraph()
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = footer.add_run('Building Excellence, Delivering Quality')
run.font.italic = True
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(128, 128, 128)

# Save document
doc.save('/home/currentsuspect/.openclaw/workspace/Tilisther_Construction_Company_Profile.docx')
print("Document created successfully!")
